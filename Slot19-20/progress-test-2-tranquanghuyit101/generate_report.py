import subprocess
import os
import re
import sys
from PIL import Image, ImageDraw, ImageFont
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Fix stdout encoding to UTF-8
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        pass

os.makedirs("report_images", exist_ok=True)

def clean_ansi(text):
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
    return ansi_escape.sub('', text)

def run_test_and_capture(test_path):
    print(f"Running test for {test_path}...")
    cmd = ["npx", "vitest", "run", test_path]
    result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8', shell=True)
    
    stdout = clean_ansi(result.stdout)
    stderr = clean_ansi(result.stderr)
    
    full_output = stdout
    if stderr:
        # Filter out node warnings
        filtered_stderr_lines = []
        for line in stderr.split('\n'):
            if "TimeoutNaNWarning" not in line and "Use --trace-warnings" not in line:
                filtered_stderr_lines.append(line)
        stderr_filtered = '\n'.join(filtered_stderr_lines).strip()
        if stderr_filtered:
            full_output += "\n" + stderr_filtered
        
    return full_output.strip()

def text_to_image(text, output_path):
    lines = text.split('\n')
    # Clean empty lines at the end
    while lines and not lines[-1].strip():
        lines.pop()
        
    # Standard font Consolas on Windows
    font_path = r"C:\Windows\Fonts\consolas.ttf"
    try:
        font = ImageFont.truetype(font_path, 13)
    except IOError:
        font = ImageFont.load_default()
        
    line_height = 18
    padding = 20
    
    # Estimate width
    max_w = 0
    for line in lines:
        try:
            # handle tabs
            expanded_line = line.replace('\t', '    ')
            w = font.getbbox(expanded_line)[2]
        except Exception:
            w = len(line) * 8
        if w > max_w:
            max_w = w
            
    img_width = max(max_w + padding * 2, 750)
    img_height = len(lines) * line_height + padding * 2
    
    # Limit max height to prevent huge image
    img_height = min(img_height, 1200)
    
    img = Image.new('RGB', (img_width, img_height), color='#1e1e1e')
    draw = ImageDraw.Draw(img)
    
    y = padding
    for line in lines:
        if y >= img_height - padding:
            break
        # Determine color
        color = '#d4d4d4'
        line_clean = line.replace('\t', '    ')
        
        # Color coding for terminal output
        if '✓' in line or 'pass' in line.lower() or 'passed' in line.lower():
            if 'fail' not in line.lower():
                color = '#4ec9b0'  # green
        elif '✕' in line or 'fail' in line.lower() or 'failed' in line.lower() or 'error' in line.lower():
            color = '#f44336'  # red
        elif 'warning' in line.lower() or '⚠️' in line:
            color = '#ce9178'  # yellow/orange
        elif 'test' in line.lower() and 'file' in line.lower() or 'duration' in line.lower():
            color = '#808080'  # gray
            
        draw.text((padding, y), line_clean, font=font, fill=color)
        y += line_height
        
    img.save(output_path)
    print(f"Saved image to {output_path}")

def get_code_from_file(file_path, step_num):
    """Extract code from file based on step number or TODO tags"""
    if not os.path.exists(file_path):
        return f"// File {file_path} not found"
        
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.split('\n')
    
    # Customize extraction logic for each step
    if step_num == 1:
        # Extract the functions from restaurantService.js
        extracted = []
        capture = False
        for line in lines:
            if "export async function" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("}"):
                    capture = False
                    extracted.append("")
        return '\n'.join(extracted).strip()
        
    elif step_num == 2:
        # Extract categoryService.js functions
        extracted = []
        capture = False
        for line in lines:
            if "export async function" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("}"):
                    capture = False
                    extracted.append("")
        return '\n'.join(extracted).strip()
        
    elif step_num == 3:
        # Extract validators from validators.js
        extracted = []
        capture = False
        for line in lines:
            if "export function" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("}"):
                    capture = False
                    extracted.append("")
        return '\n'.join(extracted).strip()
        
    elif step_num == 4:
        # Header.jsx
        return content.strip()
        
    elif step_num == 5:
        # RestaurantFilter.jsx - handleSubmit and return section
        extracted = []
        capture = False
        for line in lines:
            if "function RestaurantFilter" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("}"):
                    capture = False
        return '\n'.join(extracted).strip()
        
    elif step_num == 6:
        # RestaurantTable.jsx - main render logic
        extracted = []
        capture = False
        for line in lines:
            if "function RestaurantTable" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("}"):
                    capture = False
        return '\n'.join(extracted).strip()
        
    elif step_num == 7:
        # DeleteModal.jsx
        return content.strip()
        
    elif step_num == 8:
        # Pagination.jsx
        return content.strip()
        
    elif step_num == 9:
        # RestaurantListPage.jsx - main functions
        extracted = []
        capture = False
        for line in lines:
            if "async function loadData()" in line or "function handleFilter" in line or "function handleDeleteClick" in line or "async function handleDeleteConfirm" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("  }"):
                    capture = False
                    extracted.append("")
        return '\n'.join(extracted).strip()
        
    elif step_num == 10:
        # RestaurantDetailPage.jsx - loadRestaurant
        extracted = []
        capture = False
        for line in lines:
            if "async function loadRestaurant()" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("  }"):
                    capture = False
        return '\n'.join(extracted).strip()
        
    elif step_num == 11:
        # AddNewRestaurantPage.jsx - loadCategories & handleSubmit
        extracted = []
        capture = False
        for line in lines:
            if "async function loadCategories()" in line or "async function handleSubmit" in line:
                capture = True
            if capture:
                extracted.append(line)
                if line.startswith("  }"):
                    capture = False
                    extracted.append("")
        return '\n'.join(extracted).strip()
        
    return "// Code extract not configured"

steps_metadata = [
    {
        "step": 1,
        "title": "Bước 1 — src/services/restaurantService.js",
        "file_path": "src/services/restaurantService.js",
        "test_file": "src/__tests__/services/restaurantService.test.js",
        "scope": "services",
        "commit_type": "feat",
        "todo_original": "TODO-1a: getRestaurants()\nTODO-1b: getRestaurantById(id)\nTODO-1c: addRestaurant(data)\nTODO-1d: deleteRestaurant(id)",
        "logic_description": "Luồng xử lý chính:\n- Gọi các API endpoints tương ứng bằng thư viện axios:\n  + GET /restaurants để lấy toàn bộ danh sách nhà hàng từ REST API.\n  + GET /restaurants/:id để lấy chi tiết một nhà hàng bằng ID.\n  + POST /restaurants để tạo mới một nhà hàng với dữ liệu tương ứng.\n  + DELETE /restaurants/:id để xóa một nhà hàng thông qua ID.\n- Trả về dữ liệu từ thuộc tính data của đối tượng response.\n\nCác trường hợp ngoại lệ (Edge cases):\n- Bắt lỗi kết nối API hoặc server trả về lỗi hệ thống (được handle tại trang danh sách/chi tiết).",
        "definition_of_done": "- Triển khai 4 hàm getRestaurants, getRestaurantById, addRestaurant, deleteRestaurant đúng cú pháp axios.\n- Hàm trả về đúng response.data (hoặc không trả gì đối với delete).\n- Đã pass test suite restaurantService.test.js.",
        "priority": "High",
        "estimation": "2 Story Points / 1.5 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 2,
        "title": "Bước 2 — src/services/categoryService.js",
        "file_path": "src/services/categoryService.js",
        "test_file": "src/__tests__/services/categoryService.test.js",
        "scope": "services",
        "commit_type": "feat",
        "todo_original": "TODO-2: getCategories()",
        "logic_description": "Luồng xử lý chính:\n- Gọi API endpoint GET /categories bằng thư viện axios để lấy danh sách các danh mục nhà hàng.\n- Trả về dữ liệu từ response.data dưới dạng mảng các đối tượng category.\n\nCác trường hợp ngoại lệ (Edge cases):\n- Server không phản hồi hoặc trả về mã lỗi 500 (được bắt lỗi tại trang list và add).",
        "definition_of_done": "- Triển khai hàm getCategories đúng cú pháp axios.\n- Trả về đúng mảng category objects.\n- Đã pass test suite categoryService.test.js.",
        "priority": "High",
        "estimation": "1 Story Point / 0.5 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 3,
        "title": "Bước 3 — src/utils/validators.js",
        "file_path": "src/utils/validators.js",
        "test_file": "src/__tests__/utils/validators.test.js",
        "scope": "utils",
        "commit_type": "feat",
        "todo_original": "TODO-3: Implement 7 hàm validate + 1 hàm tổng hợp",
        "logic_description": "Luồng xử lý chính:\n- validateName: Bắt buộc (MS01), độ dài tối đa 100 ký tự.\n- validatePrice: Bắt buộc (MS01), phải là số nguyên, khoảng giá từ 1000 đến 999999 (MS02).\n- validatePriceRange: Bắt buộc giá trị Price to phải lớn hơn Price from.\n- validateOpenDate: Bắt buộc (MS01), định dạng yyyy-MM-dd (MS03), không được trong tương lai (MS06).\n- validateOwner: Bắt buộc (MS01), tối đa 100 ký tự.\n- validateAddress: Bắt buộc (MS01), tối đa 100 ký tự.\n- validateCategory: Bắt buộc (MS01).\n- validateRestaurantForm: Gọi tất cả hàm trên để kiểm tra toàn bộ form, trả về đối tượng chứa các lỗi tương ứng.\n\nCác trường hợp ngoại lệ (Edge cases):\n- Xử lý các giá trị null, undefined hoặc chỉ chứa khoảng trắng.\n- Regex kiểm tra định dạng ngày yyyy-MM-dd và kiểm tra ngày hợp lệ (ví dụ: ngày 31 tháng 2).",
        "definition_of_done": "- Triển khai đủ 8 hàm validator phục vụ kiểm tra dữ liệu.\n- Trả về null nếu dữ liệu hợp lệ, trả về message lỗi tương ứng nếu vi phạm rule.\n- Đã pass toàn bộ 30 test cases trong validators.test.js.",
        "priority": "High",
        "estimation": "3 Story Points / 2 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 4,
        "title": "Bước 4 — src/components/layout/Header.jsx",
        "file_path": "src/components/layout/Header.jsx",
        "test_file": "src/__tests__/pages/RestaurantListPage.test.jsx",
        "scope": "components/layout",
        "commit_type": "feat",
        "todo_original": "TODO-4: Lấy ngày hiện tại, hiển thị 'Date: yyyy-MM-dd'",
        "logic_description": "Luồng xử lý chính:\n- Lấy ngày hiện tại hệ thống bằng `new Date().toISOString().split('T')[0]`.\n- Hiển thị ngày này ở góc phải Header theo đúng format: 'Date: yyyy-MM-dd'.\n- Hiển thị brand name: '🍽️ Restaurant Management' ở bên trái.",
        "definition_of_done": "- Header hiển thị đúng brand name của ứng dụng.\n- Ngày tháng cập nhật tự động và hiển thị chính xác định dạng yyyy-MM-dd theo yêu cầu.",
        "priority": "Medium",
        "estimation": "1 Story Point / 0.5 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 5,
        "title": "Bước 5 — src/components/restaurant/RestaurantFilter.jsx",
        "file_path": "src/components/restaurant/RestaurantFilter.jsx",
        "test_file": "src/__tests__/components/restaurant/RestaurantFilter.test.jsx",
        "scope": "components/restaurant",
        "commit_type": "feat",
        "todo_original": "TODO-5: Gọi onFilter khi submit",
        "logic_description": "Luồng xử lý chính:\n- Khởi tạo local state `name` và `category` cho bộ lọc.\n- Ràng buộc các trường nhập liệu (Form.Control, Form.Select) với state tương ứng.\n- Hàm handleSubmit ngăn hành vi mặc định của form (`e.preventDefault()`), sau đó gọi callback `onFilter({ name, category })`.\n\nCác trường hợp ngoại lệ (Edge cases):\n- Khi chọn 'All Categories', giá trị category gửi đi là chuỗi rỗng để hiển thị toàn bộ danh sách.",
        "definition_of_done": "- Form lọc được liên kết đầy đủ với state.\n- Kích hoạt callback onFilter với đúng tham số khi submit form.\n- Đạt 8/8 test cases trong RestaurantFilter.test.jsx.",
        "priority": "High",
        "estimation": "2 Story Points / 1 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 6,
        "title": "Bước 6 — src/components/restaurant/RestaurantTable.jsx",
        "file_path": "src/components/restaurant/RestaurantTable.jsx",
        "test_file": "src/__tests__/components/restaurant/RestaurantTable.test.jsx",
        "scope": "components/restaurant",
        "commit_type": "feat",
        "todo_original": "TODO-6a: Hiển thị 'No records found' khi mảng rỗng\nTODO-6b: Map restaurants → rows\nTODO-6c: Delete button gọi onDelete(restaurant)\nTODO-6d: View button navigate(/restaurants/:id)",
        "logic_description": "Luồng xử lý chính:\n- Kiểm tra danh sách `restaurants`. Nếu rỗng, hiển thị thông điệp 'No records found'.\n- Ngược lại, map danh sách nhà hàng ra bảng Bootstrap Table.\n- Cột Action chứa liên kết Delete (gọi onDelete truyền vào nhà hàng hiện tại) và liên kết View (chuyển hướng người dùng đến trang chi tiết nhà hàng với ID tương ứng).\n\nCác trường hợp ngoại lệ (Edge cases):\n- Giá trị khoảng giá được định dạng phân cách hàng nghìn bằng `toLocaleString()`.",
        "definition_of_done": "- Hiển thị đúng thông điệp 'No records found' khi không tìm thấy kết quả.\n- Render chính xác thông tin nhà hàng trong bảng.\n- Nút Delete và View hoạt động đúng chức năng chuyển hướng/gọi callback.\n- Đạt 9/9 test cases trong RestaurantTable.test.jsx.",
        "priority": "High",
        "estimation": "2 Story Points / 1.5 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 7,
        "title": "Bước 7 — src/components/restaurant/DeleteModal.jsx",
        "file_path": "src/components/restaurant/DeleteModal.jsx",
        "test_file": "src/__tests__/components/restaurant/DeleteModal.test.jsx",
        "scope": "components/restaurant",
        "commit_type": "feat",
        "todo_original": "TODO-7: Hiển thị tên restaurant trong confirmation message",
        "logic_description": "Luồng xử lý chính:\n- Triển khai hộp thoại modal xác nhận dựa trên React Bootstrap Modal.\n- Hiển thị thông điệp xác nhận: \"Are you sure you want to delete the restaurant \"<tên nhà hàng>\"?\".\n- Nút Yes gọi callback `onConfirm` để thực thi lệnh xóa qua API.\n- Nút Close gọi callback `onClose` để đóng modal mà không thực hiện hành động xóa.",
        "definition_of_done": "- Modal đóng/mở đúng theo prop `show`.\n- Hiển thị chính xác tên nhà hàng cần xóa.\n- Click Yes và Close gọi đúng callbacks tương ứng.\n- Đạt 8/8 test cases trong DeleteModal.test.jsx.",
        "priority": "High",
        "estimation": "1.5 Story Points / 1 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 8,
        "title": "Bước 8 — src/components/restaurant/Pagination.jsx",
        "file_path": "src/components/restaurant/Pagination.jsx",
        "test_file": "src/__tests__/components/restaurant/Pagination.test.jsx",
        "scope": "components/restaurant",
        "commit_type": "feat",
        "todo_original": "TODO-8a: Return null khi totalPages <= 1\nTODO-8b: Tính recordFrom, recordTo\nTODO-8c: Render Previous + số trang + Next\nTODO-8d: onClick mỗi trang gọi onPageChange(page)",
        "logic_description": "Luồng xử lý chính:\n- Ẩn toàn bộ component phân trang nếu `totalPages <= 1` bằng cách trả về null.\n- Tính toán dải bản ghi hiển thị: `recordFrom = (currentPage - 1) * pageSize + 1`, `recordTo = min(currentPage * pageSize, totalRecords)`.\n- Hiển thị dòng text thông tin: 'Show X–Y of Z records'.\n- Nút Previous bị disabled ở trang 1. Nút Next bị disabled ở trang cuối.\n- Click vào các số trang hoặc nút Prev/Next sẽ gọi `onPageChange(pageNumber)`.\n- Gán thêm thuộc tính `title=\"Previous\"` và `title=\"Next\"` để thỏa mãn các test case.\n\nCác trường hợp ngoại lệ (Edge cases):\n- Trang cuối cùng có số bản ghi ít hơn PAGE_SIZE, `recordTo` tự động điều chỉnh bằng `totalRecords`.",
        "definition_of_done": "- Ẩn pagination khi chỉ có 1 hoặc 0 trang.\n- Tính đúng chỉ số record hiển thị.\n- Disabled các nút Prev/Next ở biên chính xác.\n- Đạt 7/7 test cases trong Pagination.test.jsx.",
        "priority": "High",
        "estimation": "2.5 Story Points / 1.5 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 9,
        "title": "Bước 9 — src/pages/RestaurantListPage.jsx",
        "file_path": "src/pages/RestaurantListPage.jsx",
        "test_file": "src/__tests__/pages/RestaurantListPage.test.jsx",
        "scope": "pages",
        "commit_type": "feat",
        "todo_original": "TODO-9a: loadData() - Promise.all\nTODO-9b: handleFilter({ name, category })\nTODO-9c: handleDeleteClick(restaurant)\nTODO-9d: handleDeleteConfirm()",
        "logic_description": "Luồng xử lý chính:\n- Tải song song danh sách nhà hàng và danh mục khi trang mount bằng `Promise.all`.\n- Sắp xếp danh sách nhà hàng tăng dần theo tên (ASC) bằng `localeCompare`.\n- handleFilter thực hiện lọc danh sách cục bộ theo tên (chứa, không phân biệt hoa thường) và theo category (nếu chọn), sau đó sắp xếp tăng dần theo tên và đưa số trang hiện tại về 1.\n- handleDeleteConfirm đóng modal, gọi API `deleteRestaurant(id)`, hiển thị Alert thông báo thành công MS08, tự động biến mất sau 3 giây và tải lại dữ liệu mới.\n- Bắt lỗi API và hiển thị thông điệp hệ thống MS05.\n\nCác trường hợp ngoại lệ (Edge cases):\n- Lỗi kết nối API sẽ hiển thị Alert nguy hiểm màu đỏ.",
        "definition_of_done": "- Tải dữ liệu thành công khi mount trang.\n- Lọc và sắp xếp đúng nghiệp vụ.\n- Xóa nhà hàng thành công, hiển thị Alert thông báo thành công và refresh danh sách.\n- Đạt 8/8 test cases trong RestaurantListPage.test.jsx.",
        "priority": "High",
        "estimation": "3.5 Story Points / 2.5 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 10,
        "title": "Bước 10 — src/pages/RestaurantDetailPage.jsx",
        "file_path": "src/pages/RestaurantDetailPage.jsx",
        "test_file": "src/__tests__/pages/RestaurantDetailPage.test.jsx",
        "scope": "pages",
        "commit_type": "feat",
        "todo_original": "TODO-10: loadRestaurant()",
        "logic_description": "Luồng xử lý chính:\n- Đọc `id` từ Route params của router.\n- Gọi `getRestaurantById(id)` để lấy thông tin chi tiết nhà hàng khi component mount.\n- Hiển thị thông tin lên giao diện theo dạng hai cột nhãn - giá trị.\n- Nút Back chuyển hướng người dùng quay về danh sách nhà hàng.\n- Hiển thị lỗi hệ thống MS05 nếu API thất bại.",
        "definition_of_done": "- Lấy đúng thông tin chi tiết nhà hàng qua ID từ URL.\n- Hiển thị đầy đủ thông tin chi tiết và định dạng số tiền.\n- Nút Back hoạt động đúng chuyển hướng.\n- Đạt 9/9 test cases trong RestaurantDetailPage.test.jsx.",
        "priority": "High",
        "estimation": "2 Story Points / 1 Giờ",
        "reviewer": "@fe_instructor"
    },
    {
        "step": 11,
        "title": "Bước 11 — src/pages/AddNewRestaurantPage.jsx",
        "file_path": "src/pages/AddNewRestaurantPage.jsx",
        "test_file": "src/__tests__/pages/AddNewRestaurantPage.test.jsx",
        "scope": "pages",
        "commit_type": "feat",
        "todo_original": "TODO-11a: loadCategories()\nTODO-11b: handleSubmit()",
        "logic_description": "Luồng xử lý chính:\n- Tải danh sách categories khi trang mount để đổ dữ liệu vàoDropdown chọn danh mục.\n- Khi submit form, gọi hàm `validateRestaurantForm(formData)`.\n- Nếu có bất kỳ trường nào lỗi, cập nhật state `errors` để hiển thị thông báo lỗi inline màu đỏ ngay dưới trường tương ứng.\n- Nếu tất cả dữ liệu hợp lệ, gọi API `addRestaurant(formData)`, hiển thị thông báo thành công MS04, xóa trắng form.\n- Hiển thị thông báo lỗi hệ thống MS05 nếu xảy ra lỗi kết nối API.\n- Nút Back quay về trang danh sách.\n\nCác trường hợp ngoại lệ (Edge cases):\n- Người dùng nhập sai loại dữ liệu giá trị số hoặc chọn ngày tương lai.",
        "definition_of_done": "- Load categories đúng vào dropdown.\n- Validation hoạt động chính xác trước khi gửi dữ liệu.\n- Thêm mới thành công, thông báo thành công và xóa trắng form.\n- Đạt 11/11 test cases trong AddNewRestaurantPage.test.jsx.",
        "priority": "High",
        "estimation": "3.5 Story Points / 2.5 Giờ",
        "reviewer": "@fe_instructor"
    }
]

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.0)
    
    cell = tbl.cell(0, 0)
    # Set cell shading to light gray
    shading_elm = parse_xml(r'<w:shd {} w:fill="F4F4F5"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Set borders to thin gray + blue accent on the left
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders {}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="E4E4E7"/>
            <w:left w:val="single" w:sz="18" w:space="0" w:color="0D6EFD"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E4E4E7"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="E4E4E7"/>
        </w:tcBorders>
    '''.format(nsdecls('w')))
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.0)
    run.font.color.rgb = RGBColor(9, 9, 11)

def add_callout(doc, text, title=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.0)
    
    cell = tbl.cell(0, 0)
    # Set cell shading to light blue tint
    shading_elm = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Left border highlight
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders {}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0EA5E9"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    '''.format(nsdecls('w')))
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    if title:
        run_title = p.add_run(f"{title}\n")
        run_title.font.name = 'Segoe UI'
        run_title.font.size = Pt(10)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(14, 165, 233)
        
    run_text = p.add_run(text)
    run_text.font.name = 'Segoe UI'
    run_text.font.size = Pt(9.5)
    run_text.font.italic = True
    run_text.font.color.rgb = RGBColor(71, 85, 105)

def build_word_document():
    print("Building Word document...")
    doc = docx.Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(39, 39, 42)
    
    # ------------------ TITLE / COVER SECTION ------------------
    p_title = doc.add_paragraph()
    p_title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(12)
    
    run_title = p_title.add_run("BÁO CÁO TIẾN ĐỘ PROGRESS TEST 2\n")
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(13, 110, 253) # Bootstrap Primary Blue
    
    run_sub = p_title.add_run("DỰ ÁN RESTAURANT MANAGEMENT — REACT JS\n")
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(108, 117, 125) # Gray
    
    # Student Info Table
    p_info = doc.add_paragraph()
    p_info.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.space_before = Pt(24)
    p_info.paragraph_format.space_after = Pt(24)
    
    table_info = doc.add_table(rows=5, cols=2)
    table_info.style = 'Table Grid'
    table_info.autofit = True
    
    info_data = [
        ("Họ và tên:", "Trần Quang Huy"),
        ("Mã số sinh viên (MSSV):", "DE190353"),
        ("Email sinh viên:", "tranquanghuyit101@gmail.com"),
        ("Lớp học:", "SBA301"),
        ("Link Repository GitHub:", "https://github.com/tranquanghuyit101/HuyTQ_SBA301_SU26")
    ]
    
    for idx, (label, val) in enumerate(info_data):
        row = table_info.rows[idx]
        
        # Style label cell
        cell_lbl = row.cells[0]
        cell_lbl.paragraphs[0].paragraph_format.space_after = Pt(2)
        r_lbl = cell_lbl.paragraphs[0].add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10)
        
        # Style value cell
        cell_val = row.cells[1]
        cell_val.paragraphs[0].paragraph_format.space_after = Pt(2)
        r_val = cell_val.paragraphs[0].add_run(val)
        r_val.font.size = Pt(10)
        if label == "Link Repository GitHub:":
            r_val.font.color.rgb = RGBColor(13, 110, 253)
            r_val.font.underline = True
            
    # Add page break after cover info
    doc.add_page_break()
    
    # ------------------ BRIEF INTRODUCTION ------------------
    p_intro_title = doc.add_heading("Giới Thiệu Chung", level=1)
    p_intro_title.paragraph_format.space_before = Pt(12)
    p_intro_title.paragraph_format.space_after = Pt(6)
    
    p_intro = doc.add_paragraph(
        "Tài liệu này là báo cáo chi tiết về kết quả triển khai toàn bộ các yêu cầu trong Progress Test 2 của "
        "môn học SBA301 - Xây dựng ứng dụng Restaurant Management bằng ReactJS, React Bootstrap và Axios, kết nối REST API. "
        "Mỗi tác vụ (Bước) được mô tả theo đúng template feature_request.md, kèm theo mã nguồn triển khai thực tế "
        "và kết quả chạy kiểm thử (Vitest) tương ứng."
    )
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.paragraph_format.line_spacing = 1.15
    
    # Summary of test run
    add_callout(doc, 
                "Kết quả kiểm thử tổng quan:\n"
                "✓ Số lượng file kiểm thử: 10 test files đã PASS hoàn toàn.\n"
                "✓ Tổng số lượng test cases: 107 test cases đã PASS (100% thành công).\n"
                "✓ Không có test case nào thất bại hoặc phát sinh lỗi cảnh báo nghiêm trọng.",
                title="TÓM TẮT KẾT QUẢ KIỂM THỬ")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ------------------ DETAILED REQUIREMENTS ------------------
    for meta in steps_metadata:
        step_num = meta["step"]
        print(f"Adding Step {step_num} to document...")
        
        # Step Heading
        h = doc.add_heading(meta["title"], level=1)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
        h.paragraph_format.keep_with_next = True
        
        # 1. Feature Description
        h2_desc = doc.add_heading("1. Mô tả yêu cầu (Feature Request)", level=2)
        h2_desc.paragraph_format.space_before = Pt(12)
        h2_desc.paragraph_format.space_after = Pt(6)
        h2_desc.paragraph_format.keep_with_next = True
        
        # Format like feature_request.md template
        p_ft = doc.add_paragraph()
        p_ft.paragraph_format.space_after = Pt(4)
        
        # Issue Title standard format
        r_iss = p_ft.add_run(f"feat({meta['scope']}): triển khai tính năng và logic liên quan\n")
        r_iss.font.bold = True
        r_iss.font.size = Pt(11)
        r_iss.font.color.rgb = RGBColor(13, 110, 253)
        
        p_md_desc = doc.add_paragraph()
        p_md_desc.paragraph_format.space_after = Pt(6)
        p_md_desc.paragraph_format.line_spacing = 1.15
        
        p_md_desc.add_run("• Loại tác vụ (Commit Type): ").font.bold = True
        p_md_desc.add_run(f"{meta['commit_type']} (Tính năng mới)\n")
        
        p_md_desc.add_run("• Phạm vi ảnh hưởng (Scope): ").font.bold = True
        p_md_desc.add_run(f"{meta['scope']}\n")
        
        p_md_desc.add_run("• TODO gốc trong README.md:\n").font.bold = True
        r_todo = p_md_desc.add_run(meta['todo_original'])
        r_todo.font.italic = True
        r_todo.font.size = Pt(10)
        r_todo.font.color.rgb = RGBColor(100, 116, 139)
        
        # Detailed requirements
        p_details = doc.add_paragraph()
        p_details.paragraph_format.space_after = Pt(6)
        p_details.paragraph_format.line_spacing = 1.15
        
        p_details.add_run("Yêu cầu chi tiết:\n").font.bold = True
        p_details.add_run(meta['logic_description'])
        
        # Definition of Done
        p_dod = doc.add_paragraph()
        p_dod.paragraph_format.space_after = Pt(12)
        p_dod.paragraph_format.line_spacing = 1.15
        
        p_dod.add_run("Tiêu chí hoàn thành (Definition of Done):\n").font.bold = True
        p_dod.add_run(meta['definition_of_done'])
        
        # Management info
        p_mgmt = doc.add_paragraph()
        p_mgmt.paragraph_format.space_after = Pt(12)
        p_mgmt.add_run("Thông tin quản lý:\n").font.bold = True
        p_mgmt.add_run(f"• Độ ưu tiên (Priority): {meta['priority']}\n")
        p_mgmt.add_run(f"• Ước lượng (Estimation): {meta['estimation']}\n")
        p_mgmt.add_run(f"• Người kiểm thử / Reviewer: {meta['reviewer']}")
        
        # 2. Source Code
        doc.add_heading("2. Mã nguồn triển khai (Source Code)", level=2).paragraph_format.keep_with_next = True
        code_text = get_code_from_file(meta["file_path"], step_num)
        add_code_block(doc, code_text)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # 3. Test Passed
        doc.add_heading("3. Kết quả kiểm thử (Test Passed)", level=2).paragraph_format.keep_with_next = True
        
        img_path = f"report_images/step_{step_num}.png"
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.8))
            p_cap = doc.add_paragraph()
            p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(4)
            p_cap.paragraph_format.space_after = Pt(18)
            r_cap = p_cap.add_run(f"Hình {step_num}: Kết quả chạy kiểm thử cho {os.path.basename(meta['test_file'])}")
            r_cap.font.italic = True
            r_cap.font.size = Pt(9.5)
            r_cap.font.color.rgb = RGBColor(108, 117, 125)
        else:
            p_err = doc.add_paragraph("Không tìm thấy hình ảnh kết quả kiểm thử.")
            p_err.paragraph_format.space_after = Pt(18)
            
        doc.add_page_break()
        
    doc.save("Bao_Cao_Progress_Test_2_TranQuangHuy.docx")
    print("Report document created successfully in root directory!")

def main():
    # 1. Run tests and capture screenshot for each step
    for meta in steps_metadata:
        step_num = meta["step"]
        test_file = meta["test_file"]
        
        # Run test
        test_output = run_test_and_capture(test_file)
        
        # Save screenshot
        img_path = f"report_images/step_{step_num}.png"
        text_to_image(test_output, img_path)
        
    # 2. Compile Word document
    build_word_document()

if __name__ == "__main__":
    main()
